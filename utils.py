import os
from PIL import Image
import io
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def get_stored_api_key():
    """Reads the API key from a file if it exists."""
    api_key_file = "api_key.txt"
    if os.path.exists(api_key_file):
        with open(api_key_file, "r") as file:
            return file.read().strip()
    return None

def save_api_key_to_file(api_key):
    """Saves the API key to a file."""
    with open("api_key.txt", "w") as file:
        file.write(api_key)

def convert_image_for_word(img):
    if img.mode != 'RGB':
        img = img.convert('RGB')
    
    max_width = 800
    aspect_ratio = img.height / img.width
    new_height = int(max_width * aspect_ratio)
    img = img.resize((max_width, new_height), Image.LANCZOS)

    img_byte_arr = io.BytesIO()
    img.save(img_byte_arr, format='PNG')
    img_byte_arr.seek(0)

    return img_byte_arr

def format_attendance_table(df):
    df = df.copy()
    df.columns = [col.lower() for col in df.columns]
    df = df[['name', 'application_id']]
    df.insert(0, 'sr_no', range(1, len(df) + 1))
    return df

def add_table_to_document(doc, df):
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = 'Table Grid'

    header_cells = table.rows[0].cells
    headers = ['Sr. No.', 'Name', 'Application ID']
    for i, header in enumerate(headers):
        header_cells[i].text = header
        if header_cells[i].paragraphs and header_cells[i].paragraphs[0].runs:
            run = header_cells[i].paragraphs[0].runs[0]
            run.font.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0, 0, 0)

    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, value in enumerate(row):
            row_cells[i].text = str(value)
            if row_cells[i].paragraphs and row_cells[i].paragraphs[0].runs:
                run = row_cells[i].paragraphs[0].runs[0]
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(0, 0, 0)

def add_formatted_heading(doc, text, size=14, level=1, center=False):
    paragraph = doc.add_paragraph()
    if center:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)

def add_formatted_paragraph(doc, label, value):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(f"{label}: ")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)

    run = paragraph.add_run(value)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.bold = False
    run.font.color.rgb = RGBColor(0, 0, 0)
