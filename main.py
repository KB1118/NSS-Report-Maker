import os
import io
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image
from PyQt5.QtWidgets import (QApplication, QWidget, QVBoxLayout, QHBoxLayout,
                           QFileDialog, QMessageBox, QScrollArea, QFrame,
                           QSplitter, QLabel, QLineEdit, QPushButton)
from PyQt5.QtGui import QIcon, QPixmap
from PyQt5.QtCore import Qt
import sys

from utils import (get_stored_api_key, save_api_key_to_file, convert_image_for_word,
                  format_attendance_table, add_table_to_document, add_formatted_heading,
                  add_formatted_paragraph)
from api_handler import APIHandler
from gui_components import (create_api_section, create_event_details_section,
                          create_image_sections, create_attendance_section,
                          create_preview_section, apply_styles)  # Import apply_styles

class NSSReportGenerator(QWidget):
    def __init__(self):
        super().__init__()
        self.api_handler = APIHandler()
        self.setup_window()
        self.initialize_variables()
        self.create_gui()
        self.load_api_key_on_startup()

        # Load CSS styles
        self.load_styles()

    def setup_window(self):
        self.setWindowTitle("NSS Report Generator by Krish Bhatia")
        self.setGeometry(100, 100, 1200, 900)
        self.setWindowIcon(QIcon("icon.png"))

    def initialize_variables(self):
        self.images = []
        self.attendance_file = None
        self.attendance_data = None
        self.event_flyer = None

    def create_gui(self):
        # Main layout with splitter
        main_layout = QHBoxLayout(self)
        splitter = QSplitter(Qt.Horizontal)
        main_layout.addWidget(splitter)

        # Left side (input form)
        left_widget = QWidget()
        left_layout = QVBoxLayout(left_widget)

        # Create scroll area
        scroll_area = QScrollArea()
        scroll_area.setWidgetResizable(True)
        left_layout.addWidget(scroll_area)

        # Create container for scrollable content
        scroll_container = QWidget()
        self.scrollable_layout = QVBoxLayout(scroll_container)
        scroll_area.setWidget(scroll_container)

        # API Section
        self.api_frame, self.api_key_entry, save_api_button = create_api_section(self)
        save_api_button.clicked.connect(self.save_api_key)
        self.scrollable_layout.addWidget(self.api_frame)

        # Event Details Section
        (self.event_details_frame, self.title_entry, self.date_entry,
         self.time_entry, self.description_text, self.venue_entry,
         self.club_entry) = create_event_details_section()
        self.scrollable_layout.addWidget(self.event_details_frame)

        # Image Sections
        (self.images_frame, self.add_flyer_button, self.add_images_button,
         self.image_preview_frame, self.image_preview_layout) = create_image_sections()
        self.add_flyer_button.clicked.connect(self.add_flyer)
        self.add_images_button.clicked.connect(self.add_images)
        self.scrollable_layout.addWidget(self.images_frame)

        # Attendance Section
        self.attendance_frame, upload_attendance_button = create_attendance_section()
        upload_attendance_button.clicked.connect(self.upload_attendance)
        self.scrollable_layout.addWidget(self.attendance_frame)

        # Generate Report Button
        generate_report_button = QPushButton("Generate Report")
        generate_report_button.clicked.connect(self.generate_report)
        self.scrollable_layout.addWidget(generate_report_button)

        # Add left widget to splitter
        splitter.addWidget(left_widget)

        # Right side (preview)
        right_widget = QWidget()
        right_layout = QVBoxLayout(right_widget)
        self.preview_frame, self.preview_edit = create_preview_section()
        right_layout.addWidget(self.preview_frame)
        splitter.addWidget(right_widget)

        # Set initial splitter sizes
        splitter.setSizes([600, 600])

        # Connect signals for live preview
        self.connect_signals()

        # Apply styles to all widgets
        for widget in self.findChildren(QWidget):
            apply_styles(widget)

    def load_styles(self):
        print("load_styles called")  # Add this line
        try:
            with open("styles.css", "r") as f:
                print("styles.css opened successfully")  # Add this line
                self.setStyleSheet(f.read())
        except FileNotFoundError:
            QMessageBox.critical(self, "Error", "styles.css not found")
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Error loading styles: {str(e)}")

    def connect_signals(self):
        self.title_entry.textChanged.connect(self.update_preview)
        self.date_entry.textChanged.connect(self.update_preview)
        self.time_entry.textChanged.connect(self.update_preview)
        self.venue_entry.textChanged.connect(self.update_preview)
        self.club_entry.textChanged.connect(self.update_preview)
        self.description_text.textChanged.connect(self.update_preview)

    def load_api_key_on_startup(self):
        api_key = get_stored_api_key()
        if api_key:
            try:
                self.api_key_entry.setText(api_key)
                self.api_handler.initialize_model(api_key)
            except Exception as e:
                QMessageBox.critical(self, "Error", f"Error loading API key: {str(e)}")

    def save_api_key(self):
        api_key = self.api_key_entry.text().strip()
        if api_key:
            try:
                self.api_handler.initialize_model(api_key)
                save_api_key_to_file(api_key)
                QMessageBox.information(self, "Success", "API Key saved successfully!")
            except Exception as e:
                QMessageBox.critical(self, "Error", f"Invalid API Key: {str(e)}")
        else:
            QMessageBox.critical(self, "Error", "Please enter an API Key")

    def add_flyer(self):
        file, _ = QFileDialog.getOpenFileName(
            self, "Select Event Flyer", "",
            "Image files (*.png *.jpg *.jpeg *.gif *.bmp)"
        )
        if file:
            try:
                self.event_flyer = Image.open(file)
                QMessageBox.information(self, "Success", "Event flyer added successfully!")
                self.update_preview()
            except Exception as e:
                QMessageBox.critical(self, "Error", f"Error adding flyer: {str(e)}")

    def add_images(self):
        files, _ = QFileDialog.getOpenFileNames(
            self, "Select Images", "",
            "Image files (*.png *.jpg *.jpeg *.gif *.bmp)"
        )

        for file in files:
            img_frame = QFrame()
            img_layout = QVBoxLayout(img_frame)
            self.image_preview_layout.addWidget(img_frame)

            img = Image.open(file)
            img_high_res = convert_image_for_word(img)
            
            img_copy = img.copy()  # Create a copy before thumbnailing
            img_copy.thumbnail((150, 150))
            img_byte_arr = io.BytesIO()
            img_copy.save(img_byte_arr, format='PNG')
            pixmap = QPixmap()
            pixmap.loadFromData(img_byte_arr.getvalue())
            
            label = QLabel()
            label.setPixmap(pixmap)
            label.setAlignment(Qt.AlignCenter)
            img_layout.addWidget(label)

            caption_entry = QLineEdit()
            caption_entry.setText("Enter caption")
            caption_entry.textChanged.connect(self.update_preview)
            img_layout.addWidget(caption_entry)

            self.images.append({
                'path': file,
                'caption_widget': caption_entry,
                'image': img_high_res
            })
            
        self.update_preview()

    def upload_attendance(self):
        file, _ = QFileDialog.getOpenFileName(
            self, "Select Participants List File", "",
            "Excel files (*.xlsx *.xls);;CSV files (*.csv)"
        )
        self.attendance_file = file
        if self.attendance_file:
            try:
                if self.attendance_file.endswith('.csv'):
                    self.attendance_data = pd.read_csv(self.attendance_file)
                else:
                    self.attendance_data = pd.read_excel(self.attendance_file)

                required_columns = ['name', 'application_id']
                missing_columns = [col for col in required_columns if col.lower() not in
                                 [c.lower() for c in self.attendance_data.columns]]

                if missing_columns:
                    missing_columns_lower = [col.lower() for col in missing_columns]
                    raise ValueError(f"Missing required columns: {', '.join(missing_columns_lower)}")

                QMessageBox.information(self, "Success", "Participants list uploaded successfully!")
                self.update_preview()
            except Exception as e:
                self.attendance_data = None
                QMessageBox.critical(self, "Error", f"Error reading file: {str(e)}")

    def update_preview(self):
        preview_html = f"""
        <div style="font-family: 'Times New Roman'; font-size: 12pt;">
            <h1 style="text-align: center; font-size: 14pt;">Event Report</h1>
            
            <p><b>Title:</b> {self.title_entry.text()}</p>
            <p><b>Date:</b> {self.date_entry.text()}</p>
            <p><b>Time:</b> {self.time_entry.text()}</p>
            <p><b>Venue:</b> {self.venue_entry.text()}</p>
            
            <p><b>Number of Participants:</b> {len(self.attendance_data) if self.attendance_data is not None else 0}</p>
            <p><b>Name of Student Led Club:</b> {self.club_entry.text()}</p>
            
            <h2 style="font-size: 14pt;">Pre AI Summary</h2>
            <p>{self.description_text.toPlainText()}</p>
        """

        if self.images:
            preview_html += "<h2 style='font-size: 14pt; text-align: center;'>Pictures</h2>"
            for img_data in self.images:
                caption = img_data['caption_widget'].text()
                preview_html += f"<p style='text-align: center;'>[Image Placeholder]</p>"
                preview_html += f"<p style='text-align: center;'>{caption}</p>"

        if self.event_flyer:
            preview_html += "<h2 style='font-size: 14pt; text-align: center;'>Event Flyer</h2>"
            preview_html += "<p style='text-align: center;'>[Event Flyer Placeholder]</p>"

        if self.attendance_data is not None:
            preview_html += "<h2 style='font-size: 14pt;'>Participants List (Will only take reqd. columns)</h2>"
            preview_html += self.attendance_data.to_html(index=False)

        preview_html += "</div>"
        self.preview_edit.setHtml(preview_html)

    def generate_report(self):
        if not self.api_handler.model:
            QMessageBox.critical(self, "Error", "Please save your API Key first")
            return

        try:
            doc = Document()

            # Define font settings
            font_name = 'Times New Roman'
            font_size = Pt(12)

            # Apply default style
            style = doc.styles['Normal']
            style.font.name = font_name
            style.font.size = font_size

            # Add title
            add_formatted_heading(doc, 'Event Report', size=14, center=True)

            # Add event details
            add_formatted_paragraph(doc, "Title", self.title_entry.text())
            add_formatted_paragraph(doc, "Date", self.date_entry.text())
            add_formatted_paragraph(doc, "Time", self.time_entry.text())
            add_formatted_paragraph(doc, "Venue", self.venue_entry.text())

            # Add participant count
            num_participants = len(self.attendance_data) if self.attendance_data is not None else 0
            add_formatted_paragraph(doc, "Number of Participants", str(num_participants))
            add_formatted_paragraph(doc, "Name of Student Led Club", self.club_entry.text())

            # Get summary and takeaways
            original_description = self.description_text.toPlainText().strip()
            summary = self.api_handler.get_summary(original_description)
            takeaways = self.api_handler.get_takeaways(original_description)

            # Add summary
            add_formatted_heading(doc, 'Summary', 14)
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(summary)
            run.font.name = font_name
            run.font.size = font_size

            # Add takeaways
            add_formatted_heading(doc, 'Key Problem-Focused Takeaways', 14)
            for takeaway in takeaways:
                paragraph = doc.add_paragraph()
                
                # Split the takeaway into title and description
                if ':' in takeaway:
                    title, description = takeaway.split(':', 1)
                    # Add title in bold
                    run = paragraph.add_run(title + ':')
                    run.font.name = font_name
                    run.font.size = font_size
                    run.font.bold = True
                    
                    # Add description in normal font
                    run = paragraph.add_run(description)
                    run.font.name = font_name
                    run.font.size = font_size
                    run.font.bold = False
                else:
                    # If no title/description split, just add the whole takeaway
                    run = paragraph.add_run(takeaway)
                    run.font.name = font_name
                    run.font.size = font_size

            # Add images
            if self.images:
                doc.add_page_break()
                add_formatted_heading(doc, 'Pictures', 14, center=True)
                for img_data in self.images:
                    img_paragraph = doc.add_paragraph()
                    img_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    img_stream = img_data['image']
                    img_paragraph.add_run().add_picture(img_stream, width=Inches(6))

                    caption_para = doc.add_paragraph()
                    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    caption_run = caption_para.add_run(img_data['caption_widget'].text())
                    caption_run.font.name = font_name
                    caption_run.font.size = font_size

            # Add flyer
            if self.event_flyer:
                doc.add_page_break()
                add_formatted_heading(doc, 'Event Flyer', 14, center=True)
                flyer_paragraph = doc.add_paragraph()
                flyer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                flyer_stream = convert_image_for_word(self.event_flyer)
                flyer_paragraph.add_run().add_picture(flyer_stream, width=Inches(6))

            # Add attendance data
            if self.attendance_data is not None:
                doc.add_page_break()
                add_formatted_heading(doc, 'Participants List', 14)
                formatted_df = format_attendance_table(self.attendance_data)
                add_table_to_document(doc, formatted_df)

            # Save the document
            filename = f"Event Report {self.title_entry.text()}.docx".replace(" ", "_")  # Sanitize filename
            doc.save(filename)
            QMessageBox.information(self, "Success", f"Report generated successfully as {filename}")

        except Exception as e:
            QMessageBox.critical(self, "Error", f"Error generating report: {str(e)}")

if __name__ == "__main__":
    try:
        app = QApplication(sys.argv)
        report_generator = NSSReportGenerator()
        report_generator.show()
        sys.exit(app.exec_())
    except Exception as e:
        print(f"An unexpected error occurred: {str(e)}")
